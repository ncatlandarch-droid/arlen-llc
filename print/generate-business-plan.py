#!/usr/bin/env python3
"""
generate-business-plan.py
Generates Arlan LLC Comprehensive Business Plan as a .docx file.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Colour palette ──────────────────────────────────────────────
DARK_TEAL   = RGBColor(13, 79, 79)
GOLD        = RGBColor(245, 166, 35)
WHITE       = RGBColor(255, 255, 255)
BODY_BLACK  = RGBColor(34, 34, 34)
ALT_ROW     = "F8F9FA"
HEADER_HEX  = "0D4F4F"
GOLD_HEX    = "F5A623"

OUTPUT_DIR  = os.path.dirname(os.path.abspath(__file__))
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "Arlan_LLC_Business_Plan.docx")


# ── Helper functions ────────────────────────────────────────────
def set_cell_shading(cell, hex_color):
    """Apply background shading to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def fmt_cell(cell, text, bold=False, font_size=10, color=BODY_BLACK, alignment=None):
    """Format a single table cell."""
    cell.text = ""
    p = cell.paragraphs[0]
    if alignment is not None:
        p.alignment = alignment
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = "Calibri"
    run.font.color.rgb = color
    # Set spacing
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)


def add_styled_table(doc, headers, rows, bold_rows=None, col_widths=None):
    """Create a formatted table with header shading and alternating rows."""
    bold_rows = bold_rows or []
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Style header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        fmt_cell(cell, h, bold=True, font_size=10, color=WHITE)
        set_cell_shading(cell, HEADER_HEX)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        is_bold = r_idx in bold_rows
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            fmt_cell(cell, val, bold=is_bold, font_size=10)
            if r_idx % 2 == 1:
                set_cell_shading(cell, ALT_ROW)

    # Remove default borders and set thin ones
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

    # Column widths
    if col_widths:
        for r in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(r.cells):
                    r.cells[i].width = Inches(w)

    doc.add_paragraph()  # spacer
    return table


def add_heading_styled(doc, text, level=1):
    """Add a styled heading paragraph."""
    sizes = {1: 16, 2: 14, 3: 12}
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(sizes.get(level, 12))
    run.font.name = "Calibri"
    run.font.color.rgb = DARK_TEAL
    if level == 3:
        run.italic = True
    pf = p.paragraph_format
    pf.space_before = Pt(14)
    pf.space_after = Pt(6)
    return p


def add_body(doc, text):
    """Add body text paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.font.color.rgb = BODY_BLACK
    pf = p.paragraph_format
    pf.space_after = Pt(6)
    return p


def add_highlight_box(doc, text):
    """Add a gold-accented highlight box."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(8)
    pf.space_after = Pt(8)

    # Gold left border via XML
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="24" w:space="8" w:color="{GOLD_HEX}"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)

    # Shading
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="FFF8EC" w:val="clear"/>')
    pPr.append(shd)

    # Indentation
    ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="288" w:right="288"/>')
    pPr.append(ind)

    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Calibri"
    run.font.color.rgb = GOLD
    return p


def add_page_number(section):
    """Add page numbers to the footer."""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # "Prepared by Think! Ventures Foundation" line
    run1 = p.add_run("Prepared by Think! Ventures Foundation  ·  ")
    run1.font.size = Pt(8)
    run1.font.name = "Calibri"
    run1.font.color.rgb = RGBColor(128, 128, 128)

    # Page number field
    run2 = p.add_run("Page ")
    run2.font.size = Pt(8)
    run2.font.name = "Calibri"
    run2.font.color.rgb = RGBColor(128, 128, 128)

    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run3 = p.add_run()
    run3._r.append(fldChar1)

    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run4 = p.add_run()
    run4._r.append(instrText)

    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run5 = p.add_run()
    run5._r.append(fldChar2)


# ── Main document generation ───────────────────────────────────
def generate():
    doc = Document()

    # ── Page setup ──────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(0.75)
        section.right_margin  = Inches(0.75)

    section = doc.sections[0]
    add_page_number(section)

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = BODY_BLACK

    # ════════════════════════════════════════════════════════════
    # COVER PAGE
    # ════════════════════════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph()  # push content down

    # Logo text
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ARLAN")
    run.bold = True
    run.font.size = Pt(36)
    run.font.name = "Calibri"
    run.font.color.rgb = DARK_TEAL

    # Tagline
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Brilliance in Every Detail")
    run.italic = True
    run.font.size = Pt(14)
    run.font.name = "Calibri"
    run.font.color.rgb = GOLD

    doc.add_paragraph()  # spacer

    # Horizontal rule via paragraph border
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="12" w:space="1" w:color="{HEADER_HEX}"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)

    doc.add_paragraph()

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Comprehensive Business Plan")
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = "Calibri"
    run.font.color.rgb = DARK_TEAL

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Premium Exterior & Interior Home Services  |  Greensboro, NC")
    run.font.size = Pt(13)
    run.font.name = "Calibri"
    run.font.color.rgb = BODY_BLACK

    for _ in range(6):
        doc.add_paragraph()

    # Footer info
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Prepared by Think! Ventures Foundation")
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(100, 100, 100)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("501(c)(3) Nonprofit Business Incubator  |  June 2026  |  Version 3.0")
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(130, 130, 130)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # EXECUTIVE SUMMARY
    # ════════════════════════════════════════════════════════════
    add_heading_styled(doc, "Executive Summary", 1)

    add_body(doc,
        "Arlan LLC is a premium home services company based in Greensboro, NC, "
        "co-founded by Dylan Thomas and Chris Harrison. We deliver seven core services: "
        "seasonal holiday lighting, permanent LED systems, landscape lighting, window cleaning, "
        "window replacements, garage door replacements, and drone roof inspections."
    )

    add_body(doc,
        "Arlan is the flagship company of Think! Ventures Foundation, a 501(c)(3) nonprofit "
        "that builds complete digital businesses for entrepreneurs at zero cost."
    )

    add_heading_styled(doc, "Key Financial Projections", 2)

    add_styled_table(doc,
        ["Metric", "Year 1", "Year 2", "Year 3", "Year 5"],
        [
            ["Revenue",       "$83K+",  "$163K+", "$278K+", "$500K+"],
            ["Net Margin",    "35-50%", "45-55%", "48-58%", "50-62%"],
            ["Service Areas", "1",      "1-2",    "2-3",    "3+"],
            ["Team",          "2 + 1 sub", "3-4", "5-6",   "8+"],
        ],
    )

    add_highlight_box(doc,
        "$75,500+ in Digital Infrastructure — Delivered at $0 Cost\n"
        "Premium website, AI assistant, merch store, business plan, marketing materials, "
        "owner dashboard — all funded by Think! Ventures Foundation grants."
    )

    # ════════════════════════════════════════════════════════════
    # COMPANY INFORMATION
    # ════════════════════════════════════════════════════════════
    add_heading_styled(doc, "Company Information", 1)

    add_styled_table(doc,
        ["Detail", "Information"],
        [
            ["Legal Name",         "Arlan LLC"],
            ["Entity Type",        "North Carolina Limited Liability Company"],
            ["Members",            "Dylan Thomas (50%) · Chris Harrison (50%)"],
            ["Registered Agent",   "Dylan Thomas, Greensboro, NC"],
            ["Phone",              "(336) 890-0981"],
            ["Website",            "arlanpro.com"],
            ["Cooperative Partner", "Think! Ventures Foundation (10% of net profit)"],
        ],
        col_widths=[2.5, 4.5],
    )

    # ════════════════════════════════════════════════════════════
    # SERVICES & PRICING
    # ════════════════════════════════════════════════════════════
    add_heading_styled(doc, "Services & Pricing", 1)

    add_styled_table(doc,
        ["Service", "Price Range", "Margin", "Performed By"],
        [
            ["Holiday Lighting (Residential)", "$500 - $7,000",    "40-55%", "Dylan"],
            ["Holiday Lighting (Commercial)",  "$2,500 - $25,000", "40-55%", "Dylan"],
            ["Permanent LED Systems",          "$4,000 - $12,000", "45-65%", "Dylan"],
            ["Landscape Lighting",             "$2,000 - $8,000",  "40-55%", "Dylan"],
            ["Window Cleaning",                "$150 - $500",      "60-75%", "Dylan"],
            ["Window Replacements",            "$800 - $2,000",    "30-40%", "Dylan / Sub"],
            ["Garage Door Replacement",        "$1,200 - $3,000",  "15-20%*", "Nick (Sub)"],
            ["Drone Roof Inspections",         "$150 - $1,500",    "75-90%", "Dylan / Chris"],
        ],
    )

    p = doc.add_paragraph()
    run = p.add_run("*After subcontractor payment (Arlan retains 15-20% management fee)")
    run.italic = True
    run.font.size = Pt(9)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(120, 120, 120)

    # ════════════════════════════════════════════════════════════
    # YEAR 1 FINANCIAL PROJECTIONS
    # ════════════════════════════════════════════════════════════
    add_heading_styled(doc, "Year 1 Financial Projections", 1)

    add_heading_styled(doc, "Revenue by Service", 2)

    add_styled_table(doc,
        ["Service", "Est. Jobs", "Avg Revenue", "Total"],
        [
            ["Holiday Lighting",     "12", "$2,917", "$35,000"],
            ["Permanent LED",        "5",  "$3,500", "$17,500"],
            ["Landscape Lighting",   "5",  "$1,500", "$7,500"],
            ["Window Cleaning",      "30", "$200",   "$6,000"],
            ["Window Replacements",  "3",  "$1,400", "$4,200"],
            ["Garage Doors (sub)",   "4",  "$2,100", "$8,400"],
            ["Drone Inspections",    "20", "$150",   "$3,000"],
            ["Merch Sales",          "—",  "—",      "$2,000"],
            ["TOTAL",                "",   "",       "$83,600"],
        ],
        bold_rows=[8],  # TOTAL row
    )

    add_heading_styled(doc, "Projected Profit & Loss", 2)

    add_styled_table(doc,
        ["Line Item", "Amount"],
        [
            ["Gross Revenue",                      "$83,600"],
            ["Less: Materials & Supplies (25%)",    "($20,900)"],
            ["Less: Subcontractor Payments",        "($6,720)"],
            ["Less: Operating Expenses (15%)",      "($12,540)"],
            ["Net Profit",                          "$43,440"],
            ["Less: Cooperative Share (10%)",       "($4,344)"],
            ["Distributable Profit",                "$39,096"],
        ],
        bold_rows=[4, 6],  # Net Profit, Distributable Profit
        col_widths=[4.5, 2.5],
    )

    # ════════════════════════════════════════════════════════════
    # THE COOPERATIVE ADVANTAGE
    # ════════════════════════════════════════════════════════════
    add_heading_styled(doc, "The Cooperative Advantage", 1)

    add_body(doc,
        "Unlike a franchise fee that disappears, the 10% cooperative share comes back to you."
    )

    add_heading_styled(doc, "Co-op Allocation", 2)

    add_styled_table(doc,
        ["Allocation", "% of 10%", "What It Does"],
        [
            ["Think! Foundation Mission",  "40%", "Funds next business launches"],
            ["Operator Savings Fund",      "20%", "Your emergency fund / nest egg"],
            ["Annual Patronage Dividends", "20%", "Cash back at year-end"],
            ["Growth Reinvestment",        "20%", "Better tools, marketing, training"],
        ],
    )
    # ════════════════════════════════════════════════════════════
    # HOW MONEY FLOWS — PAYMENT LIFECYCLE
    # ════════════════════════════════════════════════════════════
    add_heading_styled(doc, "How Money Flows - Payment Lifecycle", 1)

    add_body(doc,
        "All revenue flows through the Arlan LLC business bank account. "
        "Members do not receive W-2 wages or salary. Instead, as LLC members, "
        "Dylan and Chris receive owner's draws (distributions) - which is standard "
        "for multi-member LLCs and avoids double taxation."
    )

    add_heading_styled(doc, "Step-by-Step: From Customer to Your Pocket", 2)

    add_styled_table(doc,
        ["Step", "What Happens", "Example ($2,500 Job)"],
        [
            ["1. Customer Pays",
             "Customer pays Arlan LLC via check, cash, Venmo, or Stripe. "
             "ALL payments go into the Arlan LLC business bank account.",
             "$2,500 deposited"],
            ["2. Pay Job Expenses",
             "Materials, supplies, and subcontractor payments come out first.",
             "-$625 materials (25%)"],
            ["3. Pay Operating Costs",
             "Insurance, gas, truck, tools - monthly fixed costs from account.",
             "-$375 overhead (15%)"],
            ["4. Calculate Net Profit",
             "Revenue minus all expenses = net profit for the period.",
             "= $1,500 net profit"],
            ["5. Set Aside Taxes",
             "Transfer 25-30% of net profit to a separate savings account "
             "for quarterly estimated tax payments.",
             "-$375 to tax savings"],
            ["6. Co-op Share (10%)",
             "10% of net profit goes to Think! Ventures cooperative pool. "
             "Paid quarterly or monthly.",
             "-$150 to co-op"],
            ["7. Owner Distributions",
             "Remaining profit split 50/50 between Dylan and Chris. "
             "Transferred from business account to personal accounts.",
             "$487.50 to Dylan / $487.50 to Chris"],
        ],
    )

    add_heading_styled(doc, "Distribution Schedule (Recommended)", 2)

    add_styled_table(doc,
        ["Option", "How It Works", "Best For"],
        [
            ["Monthly Draws (Recommended)",
             "On the 1st of each month, calculate prior month net profit. "
             "Set aside taxes + co-op, distribute remainder 50/50.",
             "Steady income, easy to budget"],
            ["Bi-Monthly Draws",
             "On the 1st and 15th, take a fixed draw amount. "
             "True-up quarterly based on actual profit.",
             "Feels like a paycheck"],
            ["Quarterly Distributions",
             "Accumulate profit for 3 months, distribute larger quarterly. "
             "Aligns with estimated tax payment dates.",
             "Larger lump sums, simpler accounting"],
        ],
    )

    add_highlight_box(doc,
        "Recommended: Monthly draws on the 1st. Review prior month's revenue and "
        "expenses in the Owner Dashboard. Set aside 30% for taxes, pay 10% co-op, "
        "split the rest 50/50."
    )

    add_heading_styled(doc, "Bank Account Structure", 2)

    add_styled_table(doc,
        ["Account", "Purpose", "Access"],
        [
            ["Arlan LLC Checking",
             "Main operating account. All revenue deposits, "
             "expense payments, and draw transfers.",
             "Dylan + Chris (both signers)"],
            ["Arlan LLC Savings (Tax)",
             "25-30% of net profit held for quarterly taxes. "
             "Used ONLY for IRS/NC estimated payments.",
             "Dylan + Chris"],
            ["Dylan Personal Account",
             "Dylan's 50% owner draws deposited monthly.",
             "Dylan only"],
            ["Chris Personal Account",
             "Chris's 50% owner draws deposited monthly.",
             "Chris only"],
        ],
    )

    add_heading_styled(doc, "Subcontractor Payments (Nick)", 2)

    add_styled_table(doc,
        ["Step", "What Happens", "Example ($2,000 Job)"],
        [
            ["1. Customer pays Arlan",
             "Full payment into Arlan LLC checking account.",
             "$2,000 to Arlan"],
            ["2. Arlan pays Nick",
             "Nick's rate (80-85%) paid to his LLC within 7 days.",
             "$1,700 to Nick's LLC"],
            ["3. Arlan keeps fee",
             "15-20% retained for lead gen, branding, admin.",
             "$300 stays in Arlan"],
            ["4. Fee splits normally",
             "The $300 flows through normal split (expenses, taxes, co-op, 50/50).",
             "~$75 each to Dylan + Chris"],
        ],
    )

    add_heading_styled(doc, "Tax Obligations (Pass-Through)", 2)

    add_body(doc,
        "As a multi-member LLC taxed as a partnership, Arlan LLC does not pay income tax. "
        "Profits pass through to each member's personal tax return. Each member is "
        "responsible for their own taxes on their share of the profit."
    )

    add_styled_table(doc,
        ["Tax", "Rate", "Due", "How to Pay"],
        [
            ["Federal Self-Employment",
             "15.3% (SS + Medicare)",
             "Quarterly",
             "IRS Direct Pay or EFTPS"],
            ["Federal Income Tax",
             "Your tax bracket (est. 22%)",
             "Quarterly",
             "IRS Direct Pay"],
            ["NC State Income Tax",
             "4.5% flat",
             "Quarterly",
             "NC DOR website"],
            ["LLC Annual Report",
             "$200 flat fee",
             "April 15/year",
             "NC Secretary of State"],
        ],
    )

    add_body(doc,
        "Quarterly tax due dates: April 15, June 15, September 15, January 15. "
        "The Owner Dashboard automatically calculates your estimated quarterly payment."
    )


    # STARTUP COSTS
    # ════════════════════════════════════════════════════════════
    add_heading_styled(doc, "Startup Costs", 1)

    add_styled_table(doc,
        ["Item", "Cost"],
        [
            ["NC LLC Filing",                "$125"],
            ["EIN",                           "$0"],
            ["General Liability Insurance",   "$360-$600/yr"],
            ["Holiday Light Inventory",       "$3,000-$5,000"],
            ["Ladders & Equipment",           "$500-$1,000"],
            ["Window Cleaning Kit",           "$300-$500"],
            ["DJI Drone + FAA Exam",          "$2,175-$3,175"],
            ["Vehicle Branding",              "$200-$300"],
            ["Marketing (first batch)",       "$100-$200"],
            ["Website & Technology",          "$0 (Think! Ventures)"],
            ["TOTAL STARTUP",                 "$6,760-$10,800"],
        ],
        bold_rows=[10],  # TOTAL row
        col_widths=[4.5, 2.5],
    )

    # ════════════════════════════════════════════════════════════
    # GROWTH STRATEGY
    # ════════════════════════════════════════════════════════════
    add_heading_styled(doc, "Growth Strategy", 1)

    add_styled_table(doc,
        ["Phase", "Timeline", "Revenue Target", "Key Milestones"],
        [
            ["1. Hustle & Master",  "Months 1-12", "$83K",  "First 5 clients, build portfolio"],
            ["2. Reputation",       "Year 2",      "$163K", "Hire help, referral program, SEO"],
            ["3. Franchise Prep",   "Year 3",      "$278K", "Complete SOPs, proven financials"],
            ["4. Franchise Launch", "Year 3-5",    "$500K+", "Clone to new markets"],
        ],
    )

    # ════════════════════════════════════════════════════════════
    # CONTACT
    # ════════════════════════════════════════════════════════════
    add_heading_styled(doc, "Contact", 1)

    add_styled_table(doc,
        ["Name", "Role", "Contact"],
        [
            ["Dylan Thomas",  "Operations & Sales",    "(336) 890-0981 · dylan@arlanpro.com"],
            ["Chris Harrison", "Technology & Strategy", "chris@thinkventures.org"],
        ],
        col_widths=[2.0, 2.0, 3.0],
    )

    # Footer links
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("arlanpro.com  ·  think-ventures.netlify.app")
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    run.font.color.rgb = DARK_TEAL

    # ── Save ────────────────────────────────────────────────────
    doc.save(OUTPUT_FILE)
    print(f"\n[OK] Business plan generated successfully!")
    print(f"File: {OUTPUT_FILE}")


if __name__ == "__main__":
    generate()
